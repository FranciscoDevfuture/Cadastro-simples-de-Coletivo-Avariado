from docx import Document
import customtkinter as ctk
from tkinter import messagebox

# Configuração inicial do customtkinter
ctk.set_appearance_mode('System')
ctk.set_default_color_theme('blue')

# Função para coletar dados dos campos de entrada
def cadastrar_usuario():
    colaborador_otu_cam = entry_colaborador.get()
    linha = entry_linha.get()
    nome = entry_nome.get()
    chapa = entry_chapa.get()
    nome_cobr = entry_nome_cobr.get()
    chapa_2 = entry_chapa_2.get()
    motivo = entry_motivo.get()
    alerta = entry_alerta.get()
    partida_hr = entry_partida_hr.get()
    partida_anterior_prefixo = entry_partida_anterior.get()
    sos_nome = entry_sos_nome.get()
    chapa_nome = entry_chapa_nome.get()
    sos_prefixo = entry_sos_prefixo.get()
    empresa_sos = entry_empresa_sos.get()

    # Verificar se todos os campos foram preenchidos
    if not all([colaborador_otu_cam, linha, nome, chapa, nome_cobr, chapa_2, motivo, alerta,
                partida_hr, partida_anterior_prefixo, sos_nome, chapa_nome, sos_prefixo, empresa_sos]):
        messagebox.showwarning('Campos Vazios', 'Por favor, preencha todos os campos.')
        return

    # Criar documento Word
    documento = Document()
    documento.add_heading('Dados da Ocorrência', level=1)

    # Adicionar informações ao documento
    documento.add_paragraph(f'Itinerário: {linha}')
    documento.add_paragraph(f'Nome: {nome}')
    documento.add_paragraph(f'Chapa: {chapa}')
    documento.add_paragraph(f'Cobrador: {nome_cobr}')
    documento.add_paragraph(f'Chapa do cobrador: {chapa_2}')
    documento.add_paragraph(f'Motivo da Avaria: {motivo}')
    documento.add_paragraph(f'Sinal de Alerta informado: {alerta}')
    documento.add_paragraph(f'Partida a ser Realizada: {partida_hr}')
    documento.add_paragraph(f'Partida anterior/Prefixo: {partida_anterior_prefixo}')
    documento.add_paragraph(f'Nome do Mecânico: {sos_nome}')
    documento.add_paragraph(f'Chapa do Mecânico: {chapa_nome}')
    documento.add_paragraph(f'Prefixo: {sos_prefixo}')
    documento.add_paragraph(f'Empresa: {empresa_sos}')

    # Salvando o documento criado
    nome_arquivo = 'cadastro_de_avariado.docx'
    documento.save(nome_arquivo)
    messagebox.showinfo('Sucesso', f'Dados salvos no arquivo: {nome_arquivo}')

# Criar a janela principal
janela = ctk.CTk()
janela.title("Ocorrência de Coletivo Avariado" +'@ By Franco')
janela.geometry("600x800")  # Tamanho da janela

# Criar e posicionar os widgets
frame = ctk.CTkFrame(janela)
frame.pack(pady=20, padx=20, fill="both", expand=True)

label_titulo = ctk.CTkLabel(frame, text="Ocorrência de Coletivo Avariado", font=("Arial", 16, "bold"))
label_titulo.pack(pady=10)

# Campos de entrada (janelas largas)
entries = [
    ("Nome e RE", "entry_colaborador"),
    ("Linha e prefixo", "entry_linha"),
    ("Nome do Operador", "entry_nome"),
    ("Chapa", "entry_chapa"),
    ("Nome do Cobrador", "entry_nome_cobr"),
    ("Chapa do Cobrador", "entry_chapa_2"),
    ("Motivo da Avaria", "entry_motivo"),
    ("Hora do Sinal de Alerta", "entry_alerta"),
    ("Partida a ser Realizada", "entry_partida_hr"),
    ("Partida Anterior (Horário e Prefixo)", "entry_partida_anterior"),
    ("Nome do Mecânico", "entry_sos_nome"),
    ("Chapa do Mecânico", "entry_chapa_nome"),
    ("Prefixo", "entry_sos_prefixo"),
    ("Empresa", "entry_empresa_sos")
]

for placeholder, var_name in entries:
    entry = ctk.CTkEntry(frame, placeholder_text=placeholder)
    entry.pack(pady=5, fill="x")
    globals()[var_name] = entry

# Botão de cadastrar
botao_cadastrar = ctk.CTkButton(frame, text="Cadastrar", command=cadastrar_usuario)
botao_cadastrar.pack(pady=20)

# Iniciar a interface gráfica
janela.mainloop()
